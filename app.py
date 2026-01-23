import base64
import os
import re
import zipfile
import secrets
from datetime import datetime
from functools import wraps
from io import BytesIO, StringIO
import pandas as pd
import pytz
import requests
from flask import Flask, render_template, request, send_file, jsonify, redirect, url_for, session
from dotenv import load_dotenv
from authlib.integrations.flask_client import OAuth

# Load environment variables from .env file
_app_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(_app_dir, '.env'))

# Optional imports for Claude integration
try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY') or secrets.token_hex(32)

# OAuth Configuration
oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=os.environ.get('GOOGLE_CLIENT_ID'),
    client_secret=os.environ.get('GOOGLE_CLIENT_SECRET'),
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
        'scope': 'openid email profile'
    }
)

# Allowed email domain
ALLOWED_EMAIL_DOMAIN = 'rzero.com'


def login_required(f):
    """Decorator to require authentication for routes"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Constants
GONG_BASE_URL = "https://us-11211.api.gong.io"
SF_TZ = pytz.timezone('America/Los_Angeles')
APP_DIR = os.path.dirname(os.path.abspath(__file__))
RUNS_DIR = os.path.join(APP_DIR, "runs")
PROMPTS_DIR = os.path.join(APP_DIR, "prompts")
os.makedirs(RUNS_DIR, exist_ok=True)
BATCH_SIZE = 10
TRANSCRIPT_BATCH_SIZE = 50
SHEET_ID = "1rOzUYCNSrxjwPI5LVUS5D7vWXRN1SIAE0gsiPIgv5l0"

# Environment variable credentials
ENV_GONG_ACCESS_KEY = os.environ.get('GONG_ACCESS_KEY', '')
ENV_GONG_SECRET_KEY = os.environ.get('GONG_SECRET_KEY', '')
ENV_ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')

def has_server_credentials():
    """Check if Gong credentials are configured in environment"""
    return bool(ENV_GONG_ACCESS_KEY and ENV_GONG_SECRET_KEY)

def get_anthropic_client():
    """Get Anthropic client if available and configured"""
    if not ANTHROPIC_AVAILABLE:
        return None
    if not ENV_ANTHROPIC_API_KEY:
        return None
    return anthropic.Anthropic(api_key=ENV_ANTHROPIC_API_KEY)

def load_prompt(prompt_type):
    """Load a prompt from the prompts directory"""
    prompt_files = {
        'objection': 'objection-analysis.md',
        'partnership': 'partnership-opportunities.md',
        'feedback': 'product-feedback.md',
        'needs': 'customer-needs.md'
    }

    filename = prompt_files.get(prompt_type)
    if not filename:
        return None

    filepath = os.path.join(PROMPTS_DIR, filename)
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        return None

# Product precedence order
PRODUCT_PRECEDENCE = [
    "eaas and savings measurement",
    "odcv", 
    "secure air",
    "occupancy analytics",
    "iaq monitoring"
]

# Product abbreviations for ranking
PRODUCT_ABBREVIATIONS = {
    "secure air": "SA",
    "odcv": "ODCV", 
    "occupancy analytics": "Occ",
    "iaq monitoring": "IAQ",
    "eaas and savings measurement": "EaaS"
}

# Global variables for Google Sheets data
PRODUCT_MAPPINGS = {}
TRACKER_MAPPINGS = {}
TRACKER_TO_PRODUCT_MAPPINGS = {}
CALL_ID_TO_ACCOUNT_NAME = {}
ACCOUNT_NAME_MAPPINGS = {}
OWNER_ACCOUNT_NAMES = set()
TARGET_DOMAINS = set()
TENANT_DOMAINS = set()
INTERNAL_DOMAINS = set()
INTERNAL_SPEAKERS = set()
EXCLUDED_DOMAINS = set()
EXCLUDED_ACCOUNT_NAMES = set()
ALWAYS_INCLUDE_DOMAINS = {}

def natural_sort_key(filename):
    """Helper function for natural sorting of filenames with numbers"""
    parts = re.split(r'(\d+)', filename)
    return [int(part) if part.isdigit() else part.lower() for part in parts]

# Add natural sort filter to Jinja2
@app.template_filter('natural_sort')
def natural_sort_filter(filenames):
    return sorted(filenames, key=natural_sort_key)

def load_csv_from_sheet(gid):
    url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=csv&gid={gid}"
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            return pd.read_csv(StringIO(response.text))
    except:
        pass
    return pd.DataFrame()

def normalize_domain(url):
    if not url or url.lower() in ["n/a", "unknown", ""]:
        return "unknown"
    try:
        domain = re.sub(r'^https?://', '', str(url).lower())
        domain = re.sub(r'^www\.', '', domain)
        return domain.split('/')[0].strip() or "unknown"
    except:
        return "unknown"

def get_email_domain(email):
    if not email or "@" not in email:
        return ""
    return email.split("@")[-1].strip().lower()

def get_field(data, key, default=""):
    if not isinstance(data, dict):
        return default
    return next((v if v is not None else default for k, v in data.items() if k.lower() == key.lower()), default)

def extract_field_values(context, field_name, object_type=None):
    values = []
    for ctx in context or []:
        for obj in ctx.get("objects", []):
            if object_type and get_field(obj, "objectType", "").lower() != object_type.lower():
                continue
            if field_name.lower() == "objectid":
                if value := get_field(obj, "objectId", ""):
                    values.append(str(value))
            else:
                for field in obj.get("fields", []):
                    if isinstance(field, dict) and get_field(field, "name", "").lower() == field_name.lower():
                        if value := get_field(field, "value", ""):
                            values.append(str(value))
    return values

# Load all Google Sheets data
def initialize_data():
    global PRODUCT_MAPPINGS, TRACKER_MAPPINGS, TRACKER_TO_PRODUCT_MAPPINGS
    global CALL_ID_TO_ACCOUNT_NAME, ACCOUNT_NAME_MAPPINGS
    global OWNER_ACCOUNT_NAMES, TARGET_DOMAINS, TENANT_DOMAINS
    global INTERNAL_DOMAINS, INTERNAL_SPEAKERS
    global EXCLUDED_DOMAINS, EXCLUDED_ACCOUNT_NAMES, ALWAYS_INCLUDE_DOMAINS
    
    # Product mappings
    df = load_csv_from_sheet(1216942066)
    if not df.empty and "Product" in df.columns and "Keyword" in df.columns:
        for _, row in df.iterrows():
            product = row.get("Product", "").lower()
            keyword = row.get("Keyword", "")
            if product and keyword:
                PRODUCT_MAPPINGS.setdefault(product, []).append(re.compile(keyword, re.IGNORECASE))
    
    # Tracker mappings
    df = load_csv_from_sheet(1601335672)
    if not df.empty and "Original Tracker" in df.columns and "Mapped Tracker" in df.columns:
        for _, row in df.iterrows():
            original = row.get("Original Tracker", "").lower()
            mapped = row.get("Mapped Tracker", "").lower()
            if original and mapped:
                TRACKER_MAPPINGS[original] = mapped
    
    # Tracker to product mappings
    df = load_csv_from_sheet(2037592660)
    if not df.empty and "Tracker" in df.columns and "Product" in df.columns:
        for _, row in df.iterrows():
            tracker = row.get("Tracker", "").lower()
            product = row.get("Product", "").lower()
            if tracker and product:
                TRACKER_TO_PRODUCT_MAPPINGS[tracker] = product
    
    # Call ID to account name
    df = load_csv_from_sheet(300481101)
    if not df.empty and "Call ID" in df.columns and "Account Name" in df.columns:
        for _, row in df.iterrows():
            call_id = str(row.get("Call ID", ""))
            account_name = row.get("Account Name", "").lower()
            if call_id and account_name:
                CALL_ID_TO_ACCOUNT_NAME[call_id] = account_name
    
    # Account name mappings
    df = load_csv_from_sheet(1023256128)
    if not df.empty and "Original Name" in df.columns and "Mapped Name" in df.columns:
        for _, row in df.iterrows():
            original = row.get("Original Name", "").lower()
            mapped = row.get("Mapped Name", "").lower()
            if original and mapped:
                ACCOUNT_NAME_MAPPINGS[original] = mapped
    
    # Owner account names
    df = load_csv_from_sheet(583478969)
    if not df.empty and "Account Name" in df.columns:
        OWNER_ACCOUNT_NAMES.update(df["Account Name"].dropna().astype(str).str.lower())
    
    # Target domains (owner domains)
    df = load_csv_from_sheet(1010248949)
    if not df.empty and "Domain" in df.columns:
        TARGET_DOMAINS.update(normalize_domain(d) for d in df["Domain"].dropna().astype(str))
    
    # Tenant domains
    df = load_csv_from_sheet(139303828)
    if not df.empty and "Domain" in df.columns:
        TENANT_DOMAINS.update(normalize_domain(d) for d in df["Domain"].dropna().astype(str))
    
    # Internal domains
    df = load_csv_from_sheet(784372544)
    if not df.empty and "Domain" in df.columns:
        INTERNAL_DOMAINS.update(df["Domain"].dropna().astype(str).str.lower())
    
    # Internal speakers
    df = load_csv_from_sheet(1402964429)
    if not df.empty and "Speaker" in df.columns:
        INTERNAL_SPEAKERS.update(df["Speaker"].dropna().astype(str).str.lower())
    
    # Excluded domains
    df = load_csv_from_sheet(463927561)
    if not df.empty and "Domain" in df.columns:
        EXCLUDED_DOMAINS.update(df["Domain"].dropna().astype(str).str.lower())
    
    # Excluded account names
    df = load_csv_from_sheet(1453423105)
    if not df.empty and "Account Name" in df.columns:
        EXCLUDED_ACCOUNT_NAMES.update(df["Account Name"].dropna().astype(str).str.lower())
    
    # Always include domains
    df = load_csv_from_sheet(1463029381)
    if not df.empty and "Domain" in df.columns and "Product" in df.columns:
        for _, row in df.iterrows():
            domain = normalize_domain(row.get("Domain", ""))
            product = row.get("Product", "").lower()
            if domain and product:
                ALWAYS_INCLUDE_DOMAINS.setdefault(domain, []).append(product)

# Gong API Client
class GongAPIClient:
    def __init__(self, access_key, secret_key):
        self.session = requests.Session()
        credentials = base64.b64encode(f"{access_key}:{secret_key}".encode()).decode()
        self.session.headers.update({"Authorization": f"Basic {credentials}"})

    def api_call(self, method, endpoint, **kwargs):
        # BUG FIX 1: Remove extra slash
        url = f"{GONG_BASE_URL}{endpoint}"  # Fixed: removed / between base URL and endpoint
        try:
            response = self.session.request(method, url, **kwargs, timeout=30)
            if response.status_code == 200:
                return response.json()
        except:
            pass
        return None

    def fetch_call_list(self, from_date, to_date):
        call_ids = []
        cursor = None
        while True:
            params = {"fromDateTime": from_date, "toDateTime": to_date}
            if cursor:
                params["cursor"] = cursor
            response = self.api_call("GET", "/v2/calls", params=params)
            if not response:
                break
            for call in response.get("calls", []):
                if call_id := call.get("id"):
                    call_ids.append(str(call_id))
            cursor = response.get("records", {}).get("cursor")
            if not cursor:
                break
        return call_ids

    def fetch_call_details(self, call_ids):
        cursor = None
        while True:
            data = {
                "filter": {"callIds": call_ids},
                "contentSelector": {
                    "exposedFields": {
                        "parties": True,
                        "content": {
                            "trackers": True,
                            "brief": True,
                            "keyPoints": True,
                            "highlights": True,
                            "outline": True,
                            "topics": True
                        }
                    },
                    "context": "Extended"
                },
                "cursor": cursor
            }
            response = self.api_call("POST", "/v2/calls/extensive", json=data)
            if not response:
                break
            for call in response.get("calls", []):
                yield call
            cursor = response.get("records", {}).get("cursor")
            if not cursor:
                break

    def fetch_transcript(self, call_ids):
        result = {}
        cursor = None
        while True:
            data = {"filter": {"callIds": call_ids}, "cursor": cursor}
            response = self.api_call("POST", "/v2/calls/transcript", json=data)
            if not response:
                break
            for t in response.get("callTranscripts", []):
                if call_id := t.get("callId"):
                    result[str(call_id)] = t.get("transcript", [])
            cursor = response.get("records", {}).get("cursor")
            if not cursor:
                break
        return result

def convert_to_sf_time(utc_time):
    if not utc_time:
        return "N/A"
    try:
        utc_time = re.sub(r'\.\d+(?=[+-]\d{2}:\d{2})', '', utc_time.replace("Z", "+00:00"))
        return datetime.fromisoformat(utc_time).astimezone(SF_TZ).strftime("%b %d, %Y")
    except:
        return "N/A"

def check_product_keywords(call, patterns):
    # Extract and flatten outline - Edit 10: Enhanced outline processing
    outline = get_field(call.get("content", {}), "outline", "")
    if isinstance(outline, list):
        outline_texts = []
        for item in outline:
            if isinstance(item, dict):
                for key, value in item.items():
                    if isinstance(value, str):
                        outline_texts.append(value)
            elif isinstance(item, str):
                outline_texts.append(item)
        outline = " ".join(outline_texts)
    elif not isinstance(outline, str):
        outline = ""
    
    # Check all content fields
    fields = [
        get_field(call.get("metaData", {}), "title", ""),
        get_field(call.get("content", {}), "brief", ""),
        outline,
        " ".join(kp.get("text", "") for kp in call.get("content", {}).get("keyPoints", [])),
        " ".join(h.get("text", "") for h in call.get("content", {}).get("highlights", []))
    ]
    
    combined = " ".join(fields).lower()
    return any(pattern.search(combined) for pattern in patterns)

def determine_products(call):
    products = []
    
    # Check content against product patterns
    for product, patterns in PRODUCT_MAPPINGS.items():
        if check_product_keywords(call, patterns):
            products.append(product)
    
    # Check trackers
    for tracker in call.get("content", {}).get("trackers", []):
        tracker_name = get_field(tracker, "name", "").lower()
        # Apply tracker mapping
        tracker_name = TRACKER_MAPPINGS.get(tracker_name, tracker_name)
        
        # Direct tracker to product mapping
        if tracker_name in TRACKER_TO_PRODUCT_MAPPINGS:
            product = TRACKER_TO_PRODUCT_MAPPINGS[tracker_name]
            if product not in products:
                products.append(product)
        
        # Check if tracker matches product patterns
        for product, patterns in PRODUCT_MAPPINGS.items():
            if any(pattern.search(tracker_name) for pattern in patterns):
                if product not in products:
                    products.append(product)
    
    return products

def resolve_account_name(call):
    context = call.get("context", [])
    call_id = get_field(call.get("metaData", {}), "id", "")
    call_id_clean = call_id.lstrip("'")
    
    # Check override first
    if call_id_clean in CALL_ID_TO_ACCOUNT_NAME:
        return CALL_ID_TO_ACCOUNT_NAME[call_id_clean]
    
    # Get from context
    account_name = (extract_field_values(context, "name", "account") or [""])[0].lower()
    
    # Apply mapping
    account_name = ACCOUNT_NAME_MAPPINGS.get(account_name, account_name)
    
    # If no name, try website domain
    if not account_name:
        website = (extract_field_values(context, "website", "account") or [""])[0]
        if website:
            account_name = normalize_domain(website)
    
    # If still no name, infer from email domains
    if not account_name:
        email_domains = []
        for party in call.get("parties", []):
            if email := get_field(party, "emailAddress", ""):
                domain = get_email_domain(email)
                if domain and domain not in INTERNAL_DOMAINS and domain not in EXCLUDED_DOMAINS:
                    email_domains.append(domain)
        if email_domains:
            # Most common domain
            account_name = max(set(email_domains), key=email_domains.count)
    
    return account_name or "unknown"

def determine_org_type(account_name, account_website):
    normalized_domain = normalize_domain(account_website)
    
    if account_name in OWNER_ACCOUNT_NAMES or normalized_domain in TARGET_DOMAINS:
        return "owner"
    elif normalized_domain in TENANT_DOMAINS:
        return "tenant"
    else:
        return "tenant"  # Default

def should_include_call(call_info, selected_products):
    account_name = call_info["account_name"].lower()
    account_domain = normalize_domain(call_info["account_website"])
    selected_lower = [p.lower() for p in selected_products]
    
    # Check exclusions
    if account_name in EXCLUDED_ACCOUNT_NAMES:
        return False
    
    # Check email domains for exclusions
    for party in call_info["parties"]:
        if email := get_field(party, "emailAddress", ""):
            if get_email_domain(email) in EXCLUDED_DOMAINS:
                return False
    
    # Check if call has selected products
    call_products = [p.lower() for p in call_info["products"]]
    if any(p in selected_lower for p in call_products):
        return True
    
    # Check always include domains
    if account_domain in ALWAYS_INCLUDE_DOMAINS:
        domain_products = ALWAYS_INCLUDE_DOMAINS[account_domain]
        if any(p in selected_lower for p in domain_products):
            return True
    
    return False

def is_internal_speaker(party):
    name = get_field(party, "name", "").lower()
    email = get_field(party, "emailAddress", "")
    
    # Check by name
    if name in INTERNAL_SPEAKERS:
        return True
    
    # Check by email domain - Edit 9: Subdomain matching
    if email:
        domain = get_email_domain(email)
        # Exact match
        if domain in INTERNAL_DOMAINS:
            return True
        # Subdomain match
        if any(domain.endswith("." + d) for d in INTERNAL_DOMAINS):
            return True
    
    return False

def format_transcript(call_data, transcript_data, product=None):
    # Build speaker lookup
    speakers = {}
    speaker_lines = []
    
    for party in call_data["parties"]:
        speaker_id = get_field(party, "speakerId", "")
        if not speaker_id:
            continue
        
        name = get_field(party, "name", "Unknown")
        title = get_field(party, "title", "")
        affiliation = "I" if is_internal_speaker(party) else "E"
        
        speakers[speaker_id] = {
            "first_name": name.split()[0] if name and " " in name else name or "Unknown",
            "affiliation": affiliation
        }
        
        line = f"{name} [{affiliation}]"
        if title:
            line += f": {title}"
        speaker_lines.append(line)
    
    # Get EaaS patterns if processing EaaS product
    eaas_patterns = []
    if product and product.lower() == "eaas and savings measurement":
        eaas_patterns = PRODUCT_MAPPINGS.get("eaas and savings measurement", [])
    
    # Group consecutive sentences from same speaker
    transcript_lines = []
    current_speaker = None
    current_sentences = []
    current_time_ms = 0
    
    for mono in transcript_data:
        speaker_id = mono.get("speakerId", "")
        speaker = speakers.get(speaker_id, {"first_name": "Unknown", "affiliation": "E"})
        
        for sentence in mono.get("sentences", []):
            ms = sentence.get("start", 0)
            text = sentence.get("text", "").strip()
            
            if text:
                # Edit 3: EaaS keyword tagging
                if eaas_patterns:
                    for pattern in eaas_patterns:
                        if match := pattern.search(text):
                            matched_text = match.group()
                            text = f"[ENERGY_SAVINGS: {matched_text}] {text}"
                            break
                
                # Edit 8: External speakers in ALL CAPS
                if speaker['affiliation'] != "I":
                    text = text.upper()
            
            # If speaker changed or this is the first sentence
            if current_speaker != speaker_id or not current_sentences:
                # Output previous speaker's grouped sentences
                if current_sentences:
                    minutes = current_time_ms // 60000
                    seconds = (current_time_ms % 60000) // 1000
                    prev_speaker = speakers.get(current_speaker, {"first_name": "Unknown", "affiliation": "E"})
                    
                    transcript_lines.append(f"{minutes}:{seconds:02d} | {prev_speaker['first_name']} [{prev_speaker['affiliation']}]")
                    transcript_lines.append(" ".join(current_sentences))
                    transcript_lines.append("")
                
                # Start new speaker group
                current_speaker = speaker_id
                current_sentences = [text] if text else []
                current_time_ms = ms
            else:
                # Same speaker, add to current sentences
                if text:
                    current_sentences.append(text)
    
    # Don't forget the last speaker's sentences
    if current_sentences:
        minutes = current_time_ms // 60000
        seconds = (current_time_ms % 60000) // 1000
        speaker = speakers.get(current_speaker, {"first_name": "Unknown", "affiliation": "E"})
        
        transcript_lines.append(f"{minutes}:{seconds:02d} | {speaker['first_name']} [{speaker['affiliation']}]")
        transcript_lines.append(" ".join(current_sentences))
        transcript_lines.append("")
    
    return speaker_lines, transcript_lines

def assign_to_product(products, selected_products):
    # Filter precedence to only selected products, maintaining original order
    selected_lower = [p.lower() for p in selected_products]
    active_precedence = [p for p in PRODUCT_PRECEDENCE if p in selected_lower]
    
    # Find first product in active precedence that matches call's products
    for product in active_precedence:
        if product in [p.lower() for p in products]:
            return product
    return None

def calculate_ranking_score(call, product):
    """Calculate ranking score based on keyword matches - Edit 5"""
    patterns = PRODUCT_MAPPINGS.get(product.lower(), [])
    if not patterns:
        return 0
    
    # Extract and flatten outline with enhanced processing
    outline = get_field(call.get("content", {}), "outline", "")
    if isinstance(outline, list):
        outline_texts = []
        for item in outline:
            if isinstance(item, dict):
                for key, value in item.items():
                    if isinstance(value, str):
                        outline_texts.append(value)
            elif isinstance(item, str):
                outline_texts.append(item)
        outline = " ".join(outline_texts)
    elif not isinstance(outline, str):
        outline = ""
    
    # Combine all searchable text
    fields = [
        get_field(call.get("metaData", {}), "title", ""),
        get_field(call.get("content", {}), "brief", ""),
        outline,
        " ".join(kp.get("text", "") for kp in call.get("content", {}).get("keyPoints", [])),
        " ".join(h.get("text", "") for h in call.get("content", {}).get("highlights", []))
    ]
    combined_text = " ".join(fields).lower()
    
    # Count matches
    score = 0
    for pattern in patterns:
        matches = pattern.findall(combined_text)
        score += len(matches)
    
    # Owner bonus
    if call.get("org_type") == "owner":
        score *= 1.33
    
    return score

def process_calls(calls, transcripts, selected_products):
    calls_by_product = {p.lower(): [] for p in selected_products}
    summaries = []
    
    for call in calls:
        # Extract basic info
        meta = call.get("metaData", {})
        context = call.get("context", [])
        call_id = get_field(meta, "id", "")
        
        if not call_id:
            continue
        
        # Get call details
        account_name = resolve_account_name(call)
        account_website = (extract_field_values(context, "website", "account") or [""])[0]
        account_industry = (extract_field_values(context, "industry", "account") or [""])[0]
        org_type = determine_org_type(account_name, account_website)
        products = determine_products(call)
        
        # Store org_type in call for ranking calculation
        call['org_type'] = org_type
        
        call_info = {
            "call_id": f"'{call_id}",
            "title": get_field(meta, "title", ""),
            "date": convert_to_sf_time(get_field(meta, "started")),
            "account_name": account_name,
            "account_website": account_website,
            "account_industry": account_industry,
            "org_type": org_type,
            "products": products,
            "parties": call.get("parties", []),
            "summary": get_field(call.get("content", {}), "brief", ""),
            "call": call  # Store original call for topic exclusion and ranking
        }
        
        # Check if we should include
        if not should_include_call(call_info, selected_products):
            continue
        
        # Process transcript
        if transcript := transcripts.get(call_id):
            # Assign to product file using dynamic precedence
            if product := assign_to_product(call_info["products"], selected_products):
                # Check if user selected this product
                if product in [p.lower() for p in selected_products]:
                    # Format transcript with product for EaaS tagging
                    speaker_lines, transcript_lines = format_transcript(call_info, transcript, product)
                    
                    calls_by_product[product].append({
                        "call_id": call_info["call_id"],
                        "date": call_info["date"],
                        "account_name": call_info["account_name"],
                        "account_website": call_info["account_website"],
                        "account_industry": call_info["account_industry"],
                        "org_type": call_info["org_type"],
                        "products": call_info["products"],
                        "speakers": speaker_lines,
                        "transcript": transcript_lines,
                        "call": call,  # Store original call object for ranking
                        "assigned_product": product
                    })
    
    # Edit 5: Rank calls within each product
    for product, product_calls in calls_by_product.items():
        if product_calls:
            # Calculate scores
            for call_data in product_calls:
                call_data["score"] = calculate_ranking_score(call_data["call"], product)
            
            # Sort by score (descending) and assign ranks
            product_calls.sort(key=lambda x: x["score"], reverse=True)
            for i, call_data in enumerate(product_calls):
                call_data["rank"] = i + 1
    
    # Generate summaries with ranking info - FIXED: Removed duplicate for loop
    for product, product_calls in calls_by_product.items():
        for call_data in product_calls:
            summaries.append({
                "call_id": call_data["call_id"],
                "call_title": get_field(call_data["call"].get("metaData", {}), "title", ""),
                "call_date": call_data["date"],
                "product_tags": "|".join(call_data["products"]),
                "org_type": call_data["org_type"],
                "account_name": call_data["account_name"],
                "account_website": call_data["account_website"],
                "account_industry": call_data["account_industry"],
                "transcript_bucket": call_data["assigned_product"],  # Edit 6
                "call_rank": call_data["rank"],  # Edit 6
                "call_summary": get_field(call_data["call"].get("content", {}), "brief", "")
            })
    
    return calls_by_product, summaries

def generate_files(calls_by_product, summaries, start_date, end_date):
    files = []

    # Create timestamped run folder
    run_timestamp = datetime.now(SF_TZ).strftime("%Y%m%d_%H%M%S")
    run_folder = os.path.join(RUNS_DIR, f"run_{run_timestamp}_{start_date}_to_{end_date}")
    os.makedirs(run_folder, exist_ok=True)

    # Generate transcript files - UPDATED: Split into buckets of 5 instead of 10
    for product, calls in calls_by_product.items():
        if not calls:
            continue
        
        # Sort by rank (already done in process_calls)
        calls.sort(key=lambda x: x["rank"])
        
        # Split into buckets of 5 (CHANGED FROM 10)
        for bucket_idx, i in enumerate(range(0, len(calls), 5)):
            bucket_calls = calls[i:i+5]  # CHANGED FROM 10 to 5
            
            # Generate filename with abbreviation and rank
            abbrev = PRODUCT_ABBREVIATIONS.get(product, product[:3].upper())
            filename = f"{abbrev}_rank_{bucket_idx + 1}.txt"
            filepath = os.path.join(run_folder, filename)
            
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    # Header - Edit 8: Updated header
                    f.write("[I]=Internal R-Zero, [E]=External Customer (shown in ALL CAPS)\n")
                    f.write("=" * 50 + "\n")
                    f.write(f"TRANSCRIPT FILE: {product.upper()} - RANK {bucket_idx + 1}\n")
                    f.write(f"Date Range: {start_date} to {end_date}\n")
                    f.write(f"Calls in this file: {len(bucket_calls)} (ranks {bucket_calls[0]['rank']}-{bucket_calls[-1]['rank']})\n")
                    f.write(f"Generated: {datetime.now(SF_TZ).strftime('%b %d, %Y')}\n")
                    f.write("=" * 50 + "\n\n")
                    
                    # Calls
                    for j, call in enumerate(bucket_calls):
                        if j > 0:
                            f.write("\n---\n\n")
                        
                        f.write(f"CALL: {call['call_id']} (Rank #{call['rank']})\n")
                        f.write(f"DATE: {call['date']}\n")
                        f.write(f"ACCOUNT: {call['account_name']}\n")
                        f.write(f"WEBSITE: {call['account_website']}\n")
                        f.write(f"INDUSTRY: {call['account_industry']}\n")
                        f.write(f"ORG TYPE: {call['org_type']}\n")
                        f.write(f"PRODUCTS: {', '.join(call['products'])}\n\n")
                        f.write("SPEAKERS:\n")
                        for speaker in call['speakers']:
                            f.write(f"{speaker}\n")
                        f.write("---\n\n")
                        
                        for line in call['transcript']:
                            f.write(f"{line}\n")
                
                files.append((product, filename))
            except Exception as e:
                print(f"Error writing transcript file {filename}: {str(e)}")
                continue
    
    # Generate CSV - Edit 6: Include new columns
    csv_filename = f"call-summary_{start_date}_{end_date}.csv"
    csv_path = os.path.join(run_folder, csv_filename)
    
    try:
        if summaries:
            df = pd.DataFrame(summaries)
            # Ensure correct column order
            columns = [
                "call_id", "call_title", "call_date", "product_tags",
                "org_type", "account_name", "account_website", 
                "account_industry", "transcript_bucket", "call_rank", 
                "call_summary"
            ]
            df = df[columns]
            df.to_csv(csv_path, index=False)
        else:
            # Empty CSV with headers
            pd.DataFrame(columns=[
                "call_id", "call_title", "call_date", "product_tags",
                "org_type", "account_name", "account_website", 
                "account_industry", "transcript_bucket", "call_rank",
                "call_summary"
            ]).to_csv(csv_path, index=False)
        
        files.append(("summary", csv_filename))
    except Exception as e:
        print(f"Error writing CSV file: {str(e)}")

    return run_folder, files

# Initialize on startup
initialize_data()


# ============ AUTH ROUTES ============

@app.route('/login')
def login():
    """Show login page"""
    error = request.args.get('error')
    return render_template('login.html', error=error)


@app.route('/auth/login')
def auth_login():
    """Initiate Google OAuth flow"""
    # Generate nonce for security
    nonce = secrets.token_urlsafe(16)
    session['oauth_nonce'] = nonce

    # Determine redirect URI based on environment
    if os.environ.get('RAILWAY_PUBLIC_DOMAIN'):
        redirect_uri = f"https://{os.environ['RAILWAY_PUBLIC_DOMAIN']}/auth/callback"
    else:
        redirect_uri = url_for('auth_callback', _external=True)

    return google.authorize_redirect(redirect_uri, nonce=nonce)


@app.route('/auth/callback')
def auth_callback():
    """Handle Google OAuth callback"""
    try:
        token = google.authorize_access_token()
        nonce = session.pop('oauth_nonce', None)
        user_info = google.parse_id_token(token, nonce=nonce)

        email = user_info.get('email', '')

        # Check if email is from allowed domain
        if not email.endswith(f'@{ALLOWED_EMAIL_DOMAIN}'):
            return redirect(url_for('login', error=f'Access restricted to @{ALLOWED_EMAIL_DOMAIN} accounts'))

        # Store user in session
        session['user'] = {
            'email': email,
            'name': user_info.get('name', ''),
            'picture': user_info.get('picture', '')
        }

        return redirect(url_for('index'))

    except Exception as e:
        print(f"OAuth error: {str(e)}")
        return redirect(url_for('login', error='Authentication failed. Please try again.'))


@app.route('/logout')
def logout():
    """Clear session and redirect to login"""
    session.clear()
    return redirect(url_for('login'))


# ============ PROTECTED ROUTES ============

@app.route('/')
@login_required
def index():
    return render_template('index_new.html', has_server_credentials=has_server_credentials(), user=session.get('user'))

@app.route('/old')
@login_required
def index_new():
    return render_template('index_new.html', has_server_credentials=has_server_credentials(), user=session.get('user'))

def run_analysis_for_type(run_name, analysis_type):
    """Run Claude analysis for a specific type and return results"""
    client = get_anthropic_client()
    if not client:
        return None

    # Load transcripts
    transcripts = get_run_transcripts(run_name)
    if not transcripts:
        return None

    # Combine transcripts (limit to avoid token limits)
    combined_content = ""
    for t in transcripts[:5]:  # Limit to first 5 files
        combined_content += f"\n\n=== {t['filename']} ===\n{t['content'][:10000]}"

    # Analysis type names and prompts
    analysis_names = {
        'objection': 'Objection Analysis',
        'partnership': 'Partnership Opportunities',
        'feedback': 'Product Feedback',
        'needs': 'Customer Needs'
    }

    analysis_descriptions = {
        'objection': 'customer objections',
        'partnership': 'partnership opportunities',
        'feedback': 'product feedback',
        'needs': 'customer needs'
    }

    try:
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": f"""Analyze the following sales call transcripts to identify {analysis_descriptions.get(analysis_type, 'insights')}.

Context:
- [I] = Internal R-Zero speaker
- [E] = External customer speaker (shown in ALL CAPS)
- Focus on what external speakers say

CRITICAL: For every finding, you MUST include:
1. The exact quote from the transcript (verbatim, first 200 characters)
2. The call date (from the CALL: and DATE: headers)
3. The timestamp in the format MM:SS (from the transcript timestamps like "5:23 | Speaker")

Provide a structured analysis with:
1. Executive Summary (2-3 sentences)
2. Key Findings (bulleted list with quotes, call dates, and timestamps)
3. Detailed Analysis (organized by theme, with specific citations)
4. Recommendations

Example citation format:
- "QUOTE TEXT HERE..." (Call ID: 123456, Date: Jan 15, 2026, Timestamp: 5:23)

TRANSCRIPTS:
{combined_content}"""
                }
            ]
        )

        result = message.content[0].text
        return {
            'type': analysis_type,
            'name': analysis_names.get(analysis_type, analysis_type.title()),
            'content': result
        }

    except Exception as e:
        print(f"Error running {analysis_type} analysis: {str(e)}")
        return None


def save_analysis_to_file(run_folder, analysis_type, content):
    """Save analysis results to a text file in the run folder"""
    analysis_names = {
        'objection': 'Objection_Analysis',
        'partnership': 'Partnership_Opportunities',
        'feedback': 'Product_Feedback',
        'needs': 'Customer_Needs'
    }

    filename = f"{analysis_names.get(analysis_type, analysis_type)}_AI_Analysis.txt"
    filepath = os.path.join(run_folder, filename)

    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"{'=' * 60}\n")
            f.write(f"AI ANALYSIS: {analysis_names.get(analysis_type, analysis_type).replace('_', ' ').upper()}\n")
            f.write(f"Generated: {datetime.now(SF_TZ).strftime('%B %d, %Y at %I:%M %p')}\n")
            f.write(f"{'=' * 60}\n\n")
            f.write(content)
        return filename
    except Exception as e:
        print(f"Error saving analysis file: {str(e)}")
        return None


@app.route('/process', methods=['POST'])
@login_required
def process():
    # Determine which template to render based on source
    template_version = request.form.get('template', '')
    template_name = 'index_new.html' if template_version == 'new' else 'index.html'

    try:
        # Get form data - use env vars for credentials
        access_key = ENV_GONG_ACCESS_KEY
        secret_key = ENV_GONG_SECRET_KEY
        selected_products = request.form.getlist('products')
        selected_analyses = request.form.getlist('analysis_types')
        start_date = request.form.get('start_date')
        end_date = request.form.get('end_date')

        # Validate credentials
        if not access_key or not secret_key:
            return render_template(template_name,
                error="API credentials not configured. Please set GONG_ACCESS_KEY and GONG_SECRET_KEY environment variables.")

        # Validate other fields
        if not all([selected_products, start_date, end_date]):
            return render_template(template_name,
                error="Please select at least one product and specify date range")

        # Parse dates
        start_dt = datetime.strptime(start_date, '%Y-%m-%d').replace(tzinfo=pytz.UTC)
        end_dt = datetime.strptime(end_date, '%Y-%m-%d').replace(hour=23, minute=59, second=59, tzinfo=pytz.UTC)

        # Date range validation (6 months max)
        date_diff = end_dt - start_dt
        if date_diff.days > 180:
            return render_template(template_name, error="Date range cannot exceed 6 months. Please select a shorter range.")

        # Initialize API client
        client = GongAPIClient(access_key, secret_key)

        # Fetch call IDs
        call_ids = client.fetch_call_list(start_dt.isoformat(), end_dt.isoformat())
        if not call_ids:
            return render_template(template_name, error="No calls found in the selected date range")

        # Fetch transcripts in batches
        all_transcripts = {}
        for i in range(0, len(call_ids), TRANSCRIPT_BATCH_SIZE):
            batch = call_ids[i:i + TRANSCRIPT_BATCH_SIZE]
            transcripts = client.fetch_transcript(batch)
            if transcripts:
                all_transcripts.update(transcripts)

        # Fetch call details in batches
        all_calls = []
        for i in range(0, len(call_ids), BATCH_SIZE):
            batch = call_ids[i:i + BATCH_SIZE]
            for call in client.fetch_call_details(batch):
                if call:
                    all_calls.append(call)

        # Process calls
        calls_by_product, summaries = process_calls(all_calls, all_transcripts, selected_products)

        # Generate files
        run_folder, files = generate_files(calls_by_product, summaries, start_date, end_date)
        run_name = os.path.basename(run_folder)

        # Run AI analyses for selected types
        analysis_results = []
        if selected_analyses and ANTHROPIC_AVAILABLE and ENV_ANTHROPIC_API_KEY:
            for analysis_type in selected_analyses:
                result = run_analysis_for_type(run_name, analysis_type)
                if result:
                    analysis_results.append(result)
                    # Save analysis to file so it's included in ZIP
                    save_analysis_to_file(run_folder, analysis_type, result['content'])

        return render_template(template_name,
            success=True,
            files=files,
            run_name=run_name,
            total_calls=len(summaries),
            analysis_results=analysis_results
        )

    except Exception as e:
        return render_template(template_name,
            error=f"Error: {str(e)}")

@app.route('/download/<run_name>/<filename>')
@login_required
def download(run_name, filename):
    filepath = os.path.join(RUNS_DIR, run_name, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    return "File not found", 404

@app.route('/download-zip/<run_name>')
@login_required
def download_zip(run_name):
    run_path = os.path.join(RUNS_DIR, run_name)
    if not os.path.exists(run_path):
        return "Run folder not found", 404

    # Create zip in memory
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for filename in os.listdir(run_path):
            filepath = os.path.join(run_path, filename)
            if os.path.isfile(filepath):
                zf.write(filepath, filename)

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=f"{run_name}.zip"
    )

def get_run_transcripts(run_name):
    """Load all transcript files from a run folder"""
    run_path = os.path.join(RUNS_DIR, run_name)
    if not os.path.exists(run_path):
        return None

    transcripts = []
    for filename in sorted(os.listdir(run_path)):
        if filename.endswith('.txt'):
            filepath = os.path.join(run_path, filename)
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    transcripts.append({
                        'filename': filename,
                        'content': f.read()
                    })
            except:
                continue
    return transcripts


@app.route('/api/analyze', methods=['POST'])
@login_required
def api_analyze():
    """Run Claude analysis on transcripts"""
    client = get_anthropic_client()
    if not client:
        return jsonify({'error': 'Claude API not configured. Set ANTHROPIC_API_KEY environment variable.'}), 400

    data = request.get_json()
    run_name = data.get('run_name')
    analysis_type = data.get('type')

    if not run_name or not analysis_type:
        return jsonify({'error': 'Missing run_name or type parameter'}), 400

    # Load transcripts
    transcripts = get_run_transcripts(run_name)
    if not transcripts:
        return jsonify({'error': f'Run folder not found: {run_name}'}), 404

    # Load prompt template
    prompt_template = load_prompt(analysis_type)
    if not prompt_template:
        return jsonify({'error': f'Unknown analysis type: {analysis_type}'}), 400

    # Combine transcripts (limit to avoid token limits)
    combined_content = ""
    for t in transcripts[:5]:  # Limit to first 5 files
        combined_content += f"\n\n=== {t['filename']} ===\n{t['content'][:10000]}"

    # Extract the extraction prompt from the template
    analysis_names = {
        'objection': 'customer objections',
        'partnership': 'partnership opportunities',
        'feedback': 'product feedback',
        'needs': 'customer needs'
    }

    try:
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": f"""Analyze the following sales call transcripts to identify {analysis_names.get(analysis_type, 'insights')}.

Context:
- [I] = Internal R-Zero speaker
- [E] = External customer speaker (shown in ALL CAPS)
- Focus on what external speakers say

Provide a structured analysis with:
1. Executive Summary (2-3 sentences)
2. Key Findings (bulleted list)
3. Detailed Analysis (organized by theme)
4. Recommendations

TRANSCRIPTS:
{combined_content}"""
                }
            ]
        )

        result = message.content[0].text
        return jsonify({'result': result})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/chat', methods=['POST'])
@login_required
def api_chat():
    """Chat with Claude about transcripts"""
    client = get_anthropic_client()
    if not client:
        return jsonify({'error': 'Claude API not configured. Set ANTHROPIC_API_KEY environment variable.'}), 400

    data = request.get_json()
    run_name = data.get('run_name')
    message = data.get('message')
    history = data.get('history', [])

    if not run_name or not message:
        return jsonify({'error': 'Missing run_name or message parameter'}), 400

    # Load transcripts for context
    transcripts = get_run_transcripts(run_name)
    if not transcripts:
        return jsonify({'error': f'Run folder not found: {run_name}'}), 404

    # Build context from transcripts (summarized to save tokens)
    transcript_summary = f"You have access to {len(transcripts)} transcript files from Gong sales calls.\n\n"
    for t in transcripts[:3]:  # Include first 3 files as context
        transcript_summary += f"=== {t['filename']} ===\n{t['content'][:5000]}\n\n"

    # Build messages with history
    messages = []

    # Add conversation history
    for h in history[-10:]:  # Limit history to last 10 messages
        messages.append({
            "role": h.get('role', 'user'),
            "content": h.get('content', '')
        })

    # Add current message
    messages.append({
        "role": "user",
        "content": message
    })

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2048,
            system=f"""You are analyzing sales call transcripts for R-Zero, a company that provides building technology solutions.

Context about the transcripts:
- [I] = Internal R-Zero speaker
- [E] = External customer speaker (shown in ALL CAPS)
- Products include: SecureAire, EaaS/Savings Measurement, ODCV, Occupancy Analytics, IAQ Monitoring

IMPORTANT: When citing from transcripts, ALWAYS include:
1. The exact quote
2. The call date (from DATE: headers)
3. The timestamp (from timestamps like "5:23 | Speaker")

Format citations like: "quote here" (Date: Jan 15, 2026, Timestamp: 5:23)

Available transcript data:
{transcript_summary}

Answer questions about patterns, objections, customer needs, and insights from these calls. Be specific and cite examples with dates and timestamps.""",
            messages=messages
        )

        return jsonify({'response': response.content[0].text})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/download-analysis', methods=['POST'])
@login_required
def api_download_analysis():
    """Download analysis as Word document"""
    if not DOCX_AVAILABLE:
        return jsonify({'error': 'Word document generation not available. Install python-docx.'}), 400

    data = request.get_json()
    run_name = data.get('run_name', 'analysis')
    analysis_type = data.get('type', 'analysis')
    content = data.get('content', '')

    if not content:
        return jsonify({'error': 'No content to download'}), 400

    try:
        # Create Word document
        doc = Document()

        # Add title
        title_names = {
            'objection': 'Objection Analysis',
            'partnership': 'Partnership Opportunities',
            'feedback': 'Product Feedback',
            'needs': 'Customer Needs'
        }
        title = doc.add_heading(title_names.get(analysis_type, 'Analysis'), 0)

        # Add metadata
        doc.add_paragraph(f"Run: {run_name}")
        doc.add_paragraph(f"Generated: {datetime.now(SF_TZ).strftime('%B %d, %Y')}")
        doc.add_paragraph("")

        # Add content
        for line in content.split('\n'):
            if line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.strip():
                doc.add_paragraph(line)

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{analysis_type}_analysis.docx"
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/get-transcript', methods=['POST'])
@login_required
def api_get_transcript():
    """Get a specific transcript for viewing in modal"""
    data = request.get_json()
    run_name = data.get('run_name')
    call_id = data.get('call_id', '')
    date = data.get('date', '')
    timestamp = data.get('timestamp', '')

    if not run_name:
        return jsonify({'error': 'Missing run_name parameter'}), 400

    # Load transcripts
    transcripts = get_run_transcripts(run_name)
    if not transcripts:
        return jsonify({'error': f'Run folder not found: {run_name}'}), 404

    # Try to find matching transcript
    # First try by call_id if provided
    if call_id:
        for t in transcripts:
            if call_id in t['content']:
                return jsonify({'transcript': t['content'], 'filename': t['filename']})

    # Try to find by date
    if date:
        for t in transcripts:
            # Check if date appears in the transcript
            if date in t['content']:
                return jsonify({'transcript': t['content'], 'filename': t['filename']})

    # If no match found, return first transcript with a note
    if transcripts:
        return jsonify({
            'transcript': transcripts[0]['content'],
            'filename': transcripts[0]['filename'],
            'note': 'Exact call not found, showing first transcript'
        })

    return jsonify({'error': 'No transcripts found'}), 404


if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
